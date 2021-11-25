from docx import Document 

#Creación del doc
document = Document()

#título
document.add_heading('Doc python', 0) 

#añadimos párrafo
p=document.add_paragraph('Párrafo')

#estilos
p.add_run('negrita').bold =True 

p.add_run('cursiva').italic =True  

p.add_run('versalita').small_caps =True 