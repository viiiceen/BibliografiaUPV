from docx import Document 

document = Document()

document.add_heading('Bibliografía',0) 

p=document.add_paragraph('inicio')

print('Seleccione el tipo de referencia bibliográfica:')
print('Libro -> 1 ')


tipo = int(input())

if tipo == int(1): 
    print('indique primer apellido del autor')
    apellido1=input()
    apellido1.capitalize()

    print('indique segundo apellido del autor')
    apellido2=input()
    apellido2.capitalize()

    print('indique nombre del autor')
    nombre=input()
    nombre=nombre[0]

    print('indique año')
    año=input()

    print('indique título')
    titulo=input()



    p.add_run(apellido1).caps=True
    p.add_run(' ').small_caps = True

    p.add_run(apellido2).small_caps = True
    p.add_run(', ').small_caps = True

    p.add_run(nombre).bold=True
    p.add_run('(' )


    p.add_run(año)
    p.add_run(')')
    
else: 
    print('mal')

document.save('fuente.docx')